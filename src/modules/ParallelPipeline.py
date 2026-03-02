"""
Parallel Processing Pipeline for RAG - Optimized CPU/GPU Distribution

Architecture:
- CPU (Threads): Document loading (I/O bound)
- CPU (Processes): Text chunking (CPU bound, removes GIL)
- GPU (Single Worker): Embedding generation (batched inference)

Performance Features:
- Automatic VRAM detection and batch size tuning
- Bounded queues to prevent RAM blowups
- Pre-batching on CPU before GPU submission
- CSV-specific optimizations
- FP16/BF16 support for GPU

The loading, chunking, and embedding functions are exposed as
standalone helpers so the sequential pipeline can reuse them.
"""

import os
import time
import queue
import hashlib
import threading
import multiprocessing as mp
from concurrent.futures import (
    ThreadPoolExecutor, ProcessPoolExecutor, as_completed, wait, FIRST_COMPLETED
)
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
import numpy as np

try:
    import torch
    TORCH_AVAILABLE = True
except ImportError:
    TORCH_AVAILABLE = False

try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False

try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False

from src.utils.Logger import get_logger

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# Supported file extensions (shared between pipelines)
# ---------------------------------------------------------------------------
SUPPORTED_EXTENSIONS = {
    ".txt", ".pdf", ".csv", ".tsv", ".json", ".md", ".html", ".htm",
    ".docx", ".doc", ".pptx", ".xlsx", ".xls",
    ".yaml", ".yml", ".ipynb",
    ".py", ".js", ".java", ".cpp", ".c", ".cs", ".go",
    ".rs", ".ts", ".jsx", ".tsx",
}


# ============================================================================
# GPU Configuration
# ============================================================================

@dataclass
class GPUConfig:
    """GPU configuration with auto-detected VRAM settings."""
    device: str
    vram_gb: float
    batch_size: int
    pre_batch_size: int
    use_fp16: bool
    is_cuda: bool

    @staticmethod
    def detect() -> "GPUConfig":
        """Detect GPU capabilities and return tuned config."""
        if not TORCH_AVAILABLE or not torch.cuda.is_available():
            logger.info("No local GPU detected — embeddings use NVIDIA API remotely")
            return GPUConfig(
                device="cpu", vram_gb=0.0,
                batch_size=64, pre_batch_size=32,
                use_fp16=False, is_cuda=False,
            )

        props = torch.cuda.get_device_properties(0)
        vram_gb = props.total_memory / (1024 ** 3)
        logger.info(f"GPU Detected: {props.name} ({vram_gb:.1f} GB VRAM)")

        if vram_gb >= 24:
            batch_size, pre_batch = 512, 64
        elif vram_gb >= 16:
            batch_size, pre_batch = 256, 48
        elif vram_gb >= 8:
            batch_size, pre_batch = 128, 32
        elif vram_gb >= 4:
            batch_size, pre_batch = 64, 24
        else:
            batch_size, pre_batch = 32, 16

        logger.info(f"Auto-tuned: batch_size={batch_size}, pre_batch={pre_batch}, fp16=True")
        return GPUConfig(
            device="cuda:0", vram_gb=vram_gb,
            batch_size=batch_size, pre_batch_size=pre_batch,
            use_fp16=True, is_cuda=True,
        )


# ============================================================================
# Pipeline Statistics
# ============================================================================

@dataclass
class PipelineStats:
    """Tracks timing and throughput for a pipeline run."""
    documents_loaded: int = 0
    chunks_created: int = 0
    embeddings_generated: int = 0
    total_time: float = 0.0
    load_time: float = 0.0
    chunk_time: float = 0.0
    embed_time: float = 0.0


# ============================================================================
# 1. DOCUMENT LOADING  (I/O bound - ThreadPoolExecutor)
# ============================================================================

def load_document(file_path: Path) -> Optional[Tuple[str, str, str, str]]:
    """Load a single document and return its content.

    Mirrors IngestSession._load_single_file so that every file type is
    returned in the format downstream TextChunker expects:
      - CSV / TSV  -> raw text  (so _chunk_csv can parse structure)
      - PDF        -> extracted text (PyMuPDF first, then unstructured)
      - DOCX       -> extracted text (python-docx first, then unstructured)
      - Excel      -> formatted rows
      - Everything else -> UTF-8 text

    Args:
        file_path: Path object pointing to the file.

    Returns:
        (doc_id, file_name, content, file_ext_without_dot) or None on failure.
    """
    try:
        file_name = file_path.name
        file_ext = file_path.suffix.lower()
        doc_id = f"doc_{hashlib.md5(file_name.encode()).hexdigest()[:8]}"
        content = ""

        # -- PDF -------------------------------------------------------
        if file_ext == ".pdf":
            try:
                import fitz  # PyMuPDF

                doc = fitz.open(str(file_path))
                pages = []
                for page in doc:
                    text = page.get_text()
                    if text.strip():
                        pages.append(text)
                page_count = len(doc)
                doc.close()
                content = "\n\n".join(pages)
                logger.info(
                    f"Streaming PDF: {file_name} ({page_count} pages)"
                )
                logger.info(
                    f"PDF loaded: {file_name} — {page_count} pages, "
                    f"{len(content)} chars extracted"
                )
            except ImportError:
                pass

            # Fallback to unstructured
            if not content.strip():
                try:
                    from unstructured.partition.auto import partition

                    elements = partition(filename=str(file_path))
                    content = "\n".join(
                        str(el).strip() for el in elements if str(el).strip()
                    )
                except Exception as exc:
                    logger.warning(f"PDF load failed for {file_name}: {exc}")

        # -- CSV / TSV -------------------------------------------------
        elif file_ext in (".csv", ".tsv"):
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()

        # -- DOCX ------------------------------------------------------
        elif file_ext == ".docx":
            try:
                import docx as python_docx

                doc = python_docx.Document(str(file_path))
                parts = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(
                            cell.text.strip()
                            for cell in row.cells
                            if cell.text.strip()
                        )
                        if row_text:
                            parts.append(row_text)
                content = "\n\n".join(parts)
                logger.info(
                    f"Loaded DOCX with python-docx: {file_name} "
                    f"({len(parts)} sections)"
                )
            except ImportError:
                pass

            if not content.strip():
                try:
                    from unstructured.partition.docx import partition_docx

                    elements = partition_docx(filename=str(file_path))
                    content = "\n\n".join(str(el) for el in elements)
                except Exception as exc:
                    logger.warning(f"DOCX load failed for {file_name}: {exc}")

        # -- Excel -----------------------------------------------------
        elif file_ext in (".xlsx", ".xls"):
            try:
                import pandas as pd

                df = pd.read_excel(str(file_path))
                header = " | ".join(str(c) for c in df.columns)
                lines = [f"Columns: {header}"]
                for _, row in df.iterrows():
                    parts = [
                        f"{col}: {str(row[col]).strip()}"
                        for col in df.columns
                        if pd.notna(row[col]) and str(row[col]).strip()
                    ]
                    if parts:
                        lines.append(" | ".join(parts))
                content = "\n".join(lines)
            except Exception as exc:
                logger.warning(f"Excel load failed for {file_name}: {exc}")

        # -- Everything else -------------------------------------------
        else:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()

        if not content or not content.strip():
            logger.warning(f"No content extracted from {file_name}")
            return None

        return (doc_id, file_name, content, file_ext.lstrip("."))

    except Exception as exc:
        logger.error(f"Error loading {file_path}: {exc}")
        return None


def load_documents_parallel(
    doc_paths: List[Path],
    max_workers: int = 8,
) -> List[Tuple[str, str, str, str]]:
    """Load multiple documents in parallel using threads (I/O bound).

    Can be called standalone from any pipeline.

    Args:
        doc_paths: List of file paths to load.
        max_workers: Thread pool size.

    Returns:
        List of (doc_id, file_name, content, file_type) tuples.
    """
    start = time.time()
    results = []

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(load_document, p): p for p in doc_paths}

        if TQDM_AVAILABLE:
            iterator = tqdm(
                as_completed(futures), total=len(futures),
                desc="Loading docs", unit="doc",
            )
        else:
            iterator = as_completed(futures)

        for future in iterator:
            doc_data = future.result()
            if doc_data is not None:
                results.append(doc_data)

    elapsed = time.time() - start
    logger.info(f"Document loading complete: {len(results)} docs in {elapsed:.2f}s")
    return results


# ============================================================================
# 2. TEXT CHUNKING  (CPU bound - ProcessPoolExecutor)
# ============================================================================

def _chunk_worker(args: Tuple[str, str, str, str, Dict]) -> List[Dict]:
    """Process-safe worker that chunks a single document.

    Runs in a child process (ProcessPoolExecutor) so it re-imports
    everything fresh.  Returns serialisable dicts to avoid pickling issues.

    Args:
        args: (doc_id, doc_name, content, file_type, chunker_config)

    Returns:
        List of chunk dicts.
    """
    try:
        doc_id, doc_name, content, file_type, chunker_config = args

        if not content or not content.strip():
            return []

        # Force fresh imports so bytecode is never stale
        import importlib
        importlib.invalidate_caches()

        from src.modules.Chunking import TextChunker, ChunkingStrategy

        chunker = TextChunker(
            chunk_size=chunker_config.get("chunk_size", 512),
            chunk_overlap=chunker_config.get("chunk_overlap", 50),
            default_strategy=ChunkingStrategy(
                chunker_config.get("strategy", "fixed_size")
            ),
            encoding_name=chunker_config.get("encoding_name", "cl100k_base"),
            min_chunk_size=chunker_config.get("min_chunk_size", 50),
            use_db=False,
        )

        chunks = chunker.chunk_text(content, doc_id, doc_name)

        # Serialise to plain dicts for cross-process transfer
        return [
            {
                "id": c.id,
                "text": c.text,
                "document_id": c.metadata.document_id,
                "document_name": c.metadata.document_name,
                "chunk_index": c.metadata.chunk_index,
                "token_count": c.metadata.token_count,
                "strategy": c.metadata.strategy,
                "page_number": c.metadata.page_number,
                "source_url": c.metadata.source_url,
            }
            for c in chunks
        ]
    except Exception as exc:
        logger.error(f"Chunking error for {args[1]}: {exc}", exc_info=True)
        return []


def chunk_documents_parallel(
    loaded_docs: List[Tuple[str, str, str, str]],
    chunker_config: Dict,
    max_workers: Optional[int] = None,
) -> List[Dict]:
    """Chunk loaded documents in parallel using processes (CPU bound).

    Can be called standalone from any pipeline.

    Args:
        loaded_docs: List of (doc_id, doc_name, content, file_type).
        chunker_config: Dict with chunk_size, chunk_overlap, strategy, etc.
        max_workers: Number of worker processes (default: cpu_count).

    Returns:
        List of chunk dicts.
    """
    start = time.time()
    max_workers = max_workers or max(1, os.cpu_count() or 4)

    all_chunks = []
    tasks = [
        (doc_id, doc_name, content, ftype, chunker_config)
        for doc_id, doc_name, content, ftype in loaded_docs
    ]

    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        future_to_name = {
            executor.submit(_chunk_worker, t): t[1] for t in tasks
        }

        for future in as_completed(future_to_name):
            doc_name = future_to_name[future]
            try:
                chunks = future.result()
                all_chunks.extend(chunks)
                logger.info(f"Chunked {doc_name}: {len(chunks)} chunks")
            except Exception as exc:
                logger.error(f"Chunk worker error for {doc_name}: {exc}")

    elapsed = time.time() - start
    logger.info(f"Chunking complete: {len(all_chunks)} chunks in {elapsed:.2f}s")
    return all_chunks


# ============================================================================
# 3. EMBEDDING GENERATION  (GPU/API bound - single worker, batched)
# ============================================================================

def embed_chunks(
    chunks: List[Dict],
    embedding_service,
    batch_size: int = 64,
) -> Tuple[np.ndarray, List[Dict], List[str]]:
    """Generate embeddings for a list of chunk dicts.

    Uses the shared EmbeddingService (with caching + retry) so behaviour
    is identical whether called from the parallel or sequential pipeline.

    Can be called standalone from any pipeline.

    Args:
        chunks: List of chunk dicts (must have 'id' and 'text' keys).
        embedding_service: An EmbeddingService instance.
        batch_size: Texts per API call.

    Returns:
        (vectors_array, metadata_list, ids_list)
    """
    start = time.time()
    dimension = embedding_service.model.dimension

    all_vectors = []
    all_metadata = []
    all_ids = []

    for batch_start in range(0, len(chunks), batch_size):
        batch = chunks[batch_start : batch_start + batch_size]
        texts = [c["text"] for c in batch]

        results = embedding_service.embed_batch(texts, batch_size=batch_size)

        for chunk, result in zip(batch, results):
            emb = result.embedding if hasattr(result, "embedding") else result
            if emb is not None and len(emb) > 0:
                all_vectors.append(np.asarray(emb, dtype=np.float32))
                all_ids.append(chunk["id"])
                all_metadata.append({
                    "chunk_id": chunk["id"],
                    "document_id": chunk.get("document_id", ""),
                    "document_name": chunk.get("document_name", ""),
                    "chunk_index": chunk.get("chunk_index", 0),
                    "token_count": chunk.get("token_count", 0),
                    "text": chunk["text"],
                    "page_number": chunk.get("page_number"),
                    "source_url": chunk.get("source_url"),
                })

    if all_vectors:
        vectors_array = np.stack(all_vectors)
    else:
        vectors_array = np.zeros((0, dimension), dtype=np.float32)

    elapsed = time.time() - start
    logger.info(f"Embedding complete: {len(all_ids)} embeddings in {elapsed:.2f}s")
    return vectors_array, all_metadata, all_ids


# ============================================================================
# Streaming Pipeline (three concurrent stages connected by queues)
# ============================================================================

class ParallelRAGPipeline:
    """High-performance parallel RAG pipeline.

    Runs three stages concurrently via bounded queues:
      1. Loading   (ThreadPoolExecutor)  ->  doc_queue
      2. Chunking  (ProcessPoolExecutor) ->  chunk_queue
      3. Embedding (main-thread batched)

    The standalone functions above (load_document, load_documents_parallel,
    chunk_documents_parallel, embed_chunks) can also be called individually
    from the sequential pipeline when needed.
    """

    def __init__(
        self,
        documents_dir: str,
        vector_store_dir: str,
        chunk_size: int = 600,
        chunk_overlap: int = 90,
        strategy: str = "fixed_size",
        encoding_name: str = "cl100k_base",
        embedding_model: str = "BAAI/bge-m3",
        embedding_provider: str = "local",
        num_loader_threads: Optional[int] = None,
        num_chunker_processes: Optional[int] = None,
        num_embed_workers: Optional[int] = None,
        gpu_config: Optional[GPUConfig] = None,
        **embedding_kwargs,
    ):
        self.documents_dir = Path(documents_dir)
        self.vector_store_dir = Path(vector_store_dir)

        # Chunker settings (passed to child processes)
        self.chunker_config = {
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
            "strategy": strategy,
            "encoding_name": encoding_name,
            "min_chunk_size": 50,
        }

        # Concurrency
        self.num_loader_threads = num_loader_threads or 8
        self.num_chunker_processes = num_chunker_processes or max(
            1, os.cpu_count() or 4
        )
        # Concurrent embedding API calls (I/O-bound: more = faster for remote API)
        self.num_embed_workers = num_embed_workers or 6

        # GPU
        self.gpu_config = gpu_config or GPUConfig.detect()

        # Embedding service (created on the main process)
        from src.modules.Embeddings import create_embedding_service

        logger.info(f"Initializing embedding service on {self.gpu_config.device}")

        # Strip keys we pass explicitly to avoid duplicates
        filtered = {
            k: v
            for k, v in embedding_kwargs.items()
            if k not in ("model_name", "device", "use_fp16")
        }
        self.embedding_service = create_embedding_service(
            model_type=embedding_provider,
            model_name=embedding_model,
            device=self.gpu_config.device,
            use_fp16=self.gpu_config.use_fp16,
            **filtered,
        )

        # FAISS index type heuristic
        from src.modules.VectorStore import FAISSVectorStore

        doc_count = sum(
            1 for p in self.documents_dir.iterdir()
            if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
        )
        estimated_chunks = doc_count * 20
        if estimated_chunks > 1_000_000:
            index_type = "ivf_flat"
            logger.info(f"Large dataset ({estimated_chunks:,} est.) -> IVF index")
        elif estimated_chunks > 100_000:
            index_type = "hnsw"
            logger.info(f"Medium dataset ({estimated_chunks:,} est.) -> HNSW index")
        else:
            index_type = "flat"
            logger.info(
                f"Small dataset ({estimated_chunks:,} est.) "
                f"-> using Flat index (exact search)"
            )

        self.vector_store = FAISSVectorStore(
            dimension=self.embedding_service.model.dimension,
            index_type=index_type,
            store_path=str(self.vector_store_dir),
        )
        
        from src.modules.Retriever import BM25Retriever
        self.bm25_retriever = BM25Retriever(store_path=str(self.vector_store_dir))

        self.stats = PipelineStats()

        logger.info(
            f"Parallel Pipeline initialized:\n"
            f"  Loader threads:      {self.num_loader_threads}\n"
            f"  Chunker processes:   {self.num_chunker_processes}\n"
            f"  Embed workers:       {self.num_embed_workers} (concurrent API calls)\n"
            f"  GPU batch size:      {self.gpu_config.batch_size}\n"
            f"  CPU pre-batch size:  {self.gpu_config.pre_batch_size}"
        )

    # ------------------------------------------------------------------
    # Stage 1 - Document Loading (threads)
    # ------------------------------------------------------------------
    def _loading_stage(
        self, doc_paths: List[Path], output_queue: queue.Queue,
    ) -> None:
        """Load documents via threads and push to *output_queue*."""
        t0 = time.time()
        with ThreadPoolExecutor(max_workers=self.num_loader_threads) as pool:
            futures = {pool.submit(load_document, p): p for p in doc_paths}

            if TQDM_AVAILABLE:
                it = tqdm(
                    as_completed(futures), total=len(futures),
                    desc="Loading docs", unit="doc",
                )
            else:
                it = as_completed(futures)

            for future in it:
                doc_data = future.result()
                if doc_data is not None:
                    output_queue.put(doc_data)
                    self.stats.documents_loaded += 1

        output_queue.put(None)  # sentinel
        self.stats.load_time = time.time() - t0
        logger.info(
            f"Document loading complete: "
            f"{self.stats.documents_loaded} docs in {self.stats.load_time:.2f}s"
        )

    # ------------------------------------------------------------------
    # Stage 2 - Text Chunking (processes)
    # ------------------------------------------------------------------
    def _chunking_stage(
        self, input_queue: queue.Queue, output_queue: queue.Queue,
    ) -> None:
        """Pull docs from *input_queue*, chunk in processes, push batches."""
        try:
            t0 = time.time()
            batch = []
            done = False

            with ProcessPoolExecutor(
                max_workers=self.num_chunker_processes
            ) as pool:
                pending = {}

                while not done or pending:
                    # Submit new tasks from the queue
                    while (
                        not done
                        and len(pending) < self.num_chunker_processes * 2
                    ):
                        try:
                            item = input_queue.get(timeout=0.5)
                        except queue.Empty:
                            break
                        if item is None:
                            done = True
                            break
                        doc_id, doc_name, content, ftype = item
                        logger.info(
                            f"Chunking: Submitting {doc_name} "
                            f"(len={len(content)})"
                        )
                        fut = pool.submit(
                            _chunk_worker,
                            (doc_id, doc_name, content, ftype,
                             self.chunker_config),
                        )
                        pending[fut] = doc_name

                    # Harvest completed futures
                    if pending:
                        completed = [f for f in list(pending) if f.done()]
                        for fut in completed:
                            try:
                                chunks = fut.result()
                                batch.extend(chunks)
                                self.stats.chunks_created += len(chunks)
                                logger.info(
                                    f"Chunked {pending[fut]}: "
                                    f"{len(chunks)} chunks"
                                )
                            except Exception as exc:
                                logger.error(
                                    f"Chunk error for {pending[fut]}: {exc}"
                                )
                            del pending[fut]

                        # Push pre-batches to embedding queue
                        while len(batch) >= self.gpu_config.pre_batch_size:
                            output_queue.put(
                                batch[: self.gpu_config.pre_batch_size]
                            )
                            batch = batch[self.gpu_config.pre_batch_size :]

                        if not completed:
                            time.sleep(0.05)
                    else:
                        time.sleep(0.05)

                # Flush remainder
                if batch:
                    output_queue.put(batch)

            output_queue.put(None)  # sentinel
            self.stats.chunk_time = time.time() - t0
            logger.info(
                f"Chunking complete: "
                f"{self.stats.chunks_created} chunks in {self.stats.chunk_time:.2f}s"
            )
        except Exception as exc:
            logger.error(f"Chunking stage crashed: {exc}", exc_info=True)
            output_queue.put(None)

    # ------------------------------------------------------------------
    # Stage 3 - Embedding (single thread, batched API calls)
    # ------------------------------------------------------------------
    def _embedding_stage(self, input_queue: queue.Queue) -> None:
        """Concurrently embed chunk batches and add vectors to FAISS.

        Uses a ThreadPoolExecutor with ``num_embed_workers`` workers so that
        multiple remote API calls are in flight simultaneously (I/O-bound).
        FAISS buffer writes are serialised inside this single thread, which
        is the only writer, so no additional locking is needed.
        """
        t0 = time.time()
        dimension = self.embedding_service.model.dimension

        FAISS_FLUSH = 2000  # accumulate this many vectors before writing to FAISS

        # Pre-allocate write buffer
        buf      = np.empty((FAISS_FLUSH, dimension), dtype=np.float32)
        buf_meta: list = []
        buf_ids:  list = []
        buf_idx = 0

        # ── helper that runs inside each worker thread ──────────────────
        def _embed_one(chunk_batch: list) -> tuple:
            """Embed one batch; returns (chunk_batch, results). OOM-safe."""
            texts      = [c["text"] for c in chunk_batch]
            batch_size = self.gpu_config.batch_size
            for retry in range(4):
                try:
                    results = self.embedding_service.embed_batch(
                        texts, batch_size=batch_size,
                    )
                    return chunk_batch, results
                except Exception as exc:
                    if "out of memory" in str(exc).lower() and retry < 3:
                        old        = batch_size
                        batch_size = max(4, old // 2)
                        logger.warning(
                            f"OOM retry {retry+1}/3: {old} -> {batch_size}"
                        )
                        if TORCH_AVAILABLE:
                            torch.cuda.empty_cache()
                            time.sleep(0.5)
                    else:
                        raise
            return chunk_batch, []  # unreachable; satisfies type-checker

        # ── sliding-window concurrent embedding ─────────────────────────
        # Keep up to num_embed_workers API calls in flight at once.
        # As soon as any future completes we drain its results into the
        # FAISS buffer and immediately submit the next batch.
        pending: dict = {}   # {future: True}
        done_reading  = False

        with ThreadPoolExecutor(
            max_workers=self.num_embed_workers,
            thread_name_prefix="embedder",
        ) as executor:
            while True:
                # ① Refill the pool to capacity
                while not done_reading and len(pending) < self.num_embed_workers:
                    try:
                        chunk_batch = input_queue.get(timeout=0.2)
                        if chunk_batch is None:
                            done_reading = True
                            break
                        fut = executor.submit(_embed_one, chunk_batch)
                        pending[fut] = True
                    except queue.Empty:
                        break

                if not pending:
                    if done_reading:
                        break          # all batches processed
                    time.sleep(0.05)   # wait for chunker to produce more
                    continue

                # ② Wait for the first future(s) to finish
                ready, _ = wait(
                    list(pending.keys()),
                    timeout=1.0,
                    return_when=FIRST_COMPLETED,
                )

                # ③ Drain completed futures into the FAISS buffer
                for fut in ready:
                    del pending[fut]
                    chunk_batch, results = fut.result()

                    for chunk, result in zip(chunk_batch, results):
                        emb = result.embedding if hasattr(result, "embedding") else result
                        if emb is not None and len(emb) > 0:
                            buf[buf_idx] = emb
                            buf_idx += 1
                            buf_ids.append(chunk["id"])
                            buf_meta.append({
                                "chunk_id":      chunk["id"],
                                "document_id":   chunk.get("document_id", ""),
                                "document_name": chunk.get("document_name", ""),
                                "chunk_index":   chunk.get("chunk_index", 0),
                                "token_count":   chunk.get("token_count", 0),
                                "text":          chunk["text"],
                                "page_number":   chunk.get("page_number"),
                                "source_url":    chunk.get("source_url"),
                            })

                    # ④ Flush to FAISS when buffer is full
                    if buf_idx >= FAISS_FLUSH:
                        self.vector_store.add_vectors_batch(
                            buf[:buf_idx].copy(), buf_meta, buf_ids,
                            batch_size=buf_idx,
                        )
                        # Add to BM25 retriever
                        self.bm25_retriever.add_texts(
                            chunk_ids=buf_ids,
                            texts=[meta["text"] for meta in buf_meta]
                        )
                        
                        self.stats.embeddings_generated += buf_idx
                        logger.info(
                            f"Embedded {self.stats.embeddings_generated} vectors so far"
                        )
                        buf_idx  = 0
                        buf_meta = []
                        buf_ids  = []

        # ── final FAISS flush ───────────────────────────────────────────
        if buf_idx > 0:
            self.vector_store.add_vectors_batch(
                buf[:buf_idx].copy(), buf_meta, buf_ids,
                batch_size=buf_idx,
            )
            # Add to BM25 retriever
            self.bm25_retriever.add_texts(
                chunk_ids=buf_ids,
                texts=[meta["text"] for meta in buf_meta]
            )
            
            self.stats.embeddings_generated += buf_idx
            logger.info(f"Flushed final {buf_idx} embeddings")

        self.stats.embed_time = time.time() - t0
        logger.info(
            f"Embedding complete: "
            f"{self.stats.embeddings_generated} embeddings "
            f"in {self.stats.embed_time:.2f}s"
        )

    # ------------------------------------------------------------------
    # Build  (orchestrates the three-stage streaming pipeline)
    # ------------------------------------------------------------------
    def build(self) -> bool:
        """Build the RAG index using the three-stage parallel pipeline.

        Returns True on success, False on failure.
        """
        # Clear stale __pycache__ so spawned processes see latest source
        import shutil
        import importlib

        modules_dir = Path(__file__).parent
        for cache_dir in modules_dir.rglob("__pycache__"):
            shutil.rmtree(cache_dir, ignore_errors=True)
        importlib.invalidate_caches()

        pipeline_start = time.time()

        logger.info("=" * 80)
        logger.info("PARALLEL RAG PIPELINE - STARTING")
        logger.info("=" * 80)

        try:
            # Discover documents
            doc_paths = [
                p
                for p in self.documents_dir.iterdir()
                if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
            ]

            if not doc_paths:
                logger.warning(f"No documents found in {self.documents_dir}")
                return False

            logger.info(f"Found {len(doc_paths)} documents to process")

            # Queue sizing based on available RAM
            if PSUTIL_AVAILABLE:
                avail_gb = psutil.virtual_memory().available / (1024 ** 3)
                dq = min(100, max(10, int(avail_gb * 5)))
                cq = min(50, max(5, int(avail_gb * 2.5)))
                logger.info(
                    f"Queue sizes: docs={dq}, chunks={cq} "
                    f"(RAM: {avail_gb:.1f}GB available)"
                )
            else:
                dq, cq = 10, 5

            doc_queue = queue.Queue(maxsize=dq)
            chunk_queue = queue.Queue(maxsize=cq)

            # Launch stages as threads (each internally uses its own pool)
            loader = threading.Thread(
                target=self._loading_stage,
                args=(doc_paths, doc_queue),
                name="LoaderThread",
            )
            chunker = threading.Thread(
                target=self._chunking_stage,
                args=(doc_queue, chunk_queue),
                name="ChunkerThread",
            )
            embedder = threading.Thread(
                target=self._embedding_stage,
                args=(chunk_queue,),
                name="EmbedderThread",
            )

            loader.start()
            chunker.start()
            embedder.start()

            loader.join()
            logger.info("[DONE] Loading stage complete")
            chunker.join()
            logger.info("[DONE] Chunking stage complete")
            embedder.join()
            logger.info("[DONE] Embedding stage complete")

            # Persist
            self.vector_store.save()
            logger.info(
                f"[DONE] Vector store saved: "
                f"{self.vector_store.get_size()} vectors"
            )
            self.bm25_retriever.build()
            self.bm25_retriever.save()
            logger.info("[DONE] BM25 sparse index built and saved")

            self.stats.total_time = time.time() - pipeline_start

            # Summary
            docs_sec = (
                self.stats.documents_loaded / self.stats.total_time
                if self.stats.total_time > 0 else 0
            )
            chunks_sec = (
                self.stats.chunks_created / self.stats.total_time
                if self.stats.total_time > 0 else 0
            )
            emb_sec = (
                self.stats.embeddings_generated / self.stats.embed_time
                if self.stats.embed_time > 0 else 0
            )

            logger.info("=" * 80)
            logger.info("PARALLEL PIPELINE COMPLETE")
            logger.info("=" * 80)
            logger.info(
                f"\n"
                f"        Documents: {self.stats.documents_loaded}\n"
                f"        Chunks: {self.stats.chunks_created}\n"
                f"        Embeddings: {self.stats.embeddings_generated}\n"
                f"        Total time: {self.stats.total_time:.2f}s\n"
                f"        Docs/sec: {docs_sec:.2f}\n"
                f"        Chunks/sec: {chunks_sec:.2f}\n"
                f"        Embeddings/sec: {emb_sec:.2f}"
            )
            logger.info("=" * 80)

            return True

        except Exception as exc:
            logger.error(f"Pipeline error: {exc}", exc_info=True)
            return False


# ============================================================================
# Factory
# ============================================================================

def create_parallel_pipeline(
    documents_dir: str,
    vector_store_dir: str,
    **kwargs,
) -> ParallelRAGPipeline:
    """Convenience factory for ParallelRAGPipeline."""
    return ParallelRAGPipeline(
        documents_dir=documents_dir,
        vector_store_dir=vector_store_dir,
        **kwargs,
    )
