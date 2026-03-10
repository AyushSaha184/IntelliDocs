"""Session-aware document ingestion with retry, concurrency, and progress tracking.

This replaces the missing IngestSession.py with a robust implementation that:
- Processes files in parallel using ThreadPoolExecutor (I/O-bound file loading)
- Retries failed embedding API calls with exponential backoff
- Isolates per-file errors (one file failing doesn't kill the session)
- Tracks progress for status polling
- Uses robust PDF parsing (OCR detection, page streaming)
- Uses streaming CSV loading for large files

This module is ONLY used by the backend API session path.
It does NOT affect the CLI --build pipeline.
"""

import os
import time
import threading
import hashlib
import traceback
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime
import psycopg2
from concurrent.futures import ThreadPoolExecutor, as_completed
from src.modules.Chunking import TextChunker
from src.modules.Embeddings import create_embedding_service
from src.modules.VectorStore import FAISSVectorStore
from src.modules.QdrantStore import QdrantSessionStore
from src.utils.Logger import get_logger
from config.config import (
    EMBEDDING_PROVIDER,
    EMBEDDING_MODEL,
    EMBEDDING_NORMALIZE,
    EMBEDDING_DIMENSION,
    EMBEDDING_TASK_TYPE,
    HF_TOKEN,
    HF_INFERENCE_PROVIDER,
    GEMINI_API_KEY,
    NVIDIA_API_KEY,
    LM_STUDIO_BASE_URL,
    LM_STUDIO_API_KEY,
    VECTOR_BACKEND,
    postgres_connect_kwargs,
)

# Import backend-specific utilities from the consolidated Loader module
from src.modules.Loader import (
    load_pdf_robust, 
    detect_scanned_pdf,
    load_csv_streaming, 
    load_csv_full, 
    detect_csv_structure
)
from src.modules.DocumentParser import StructureAnalyzer
from src.modules.MetadataEnricher import MetadataEnricher

logger = get_logger(__name__)

# Configuration
MAX_RETRIES = 3
RETRY_BASE_DELAY = 1.0  # seconds
MAX_CONCURRENT_FILES = 4  # parallel file loading threads
EMBEDDING_BATCH_SIZE = 64
SMALL_CSV_THRESHOLD = 5 * 1024 * 1024  # 5MB — use full load below this


def _build_embedding_kwargs() -> dict:
    """Build embedding service kwargs based on configured provider."""
    kwargs = {
        "model_name": EMBEDDING_MODEL,
        "normalize_embeddings": EMBEDDING_NORMALIZE
    }

    provider = EMBEDDING_PROVIDER.lower()

    if provider in ["hf", "hf-inference", "huggingface"]:
        kwargs.update({
            "api_key": HF_TOKEN,
            "provider": HF_INFERENCE_PROVIDER
        })
    elif provider in ["gemini", "google", "google-gemini"]:
        kwargs.update({
            "api_key": GEMINI_API_KEY,
            "task_type": EMBEDDING_TASK_TYPE,
            "output_dimensionality": EMBEDDING_DIMENSION
        })
    elif provider in ["nvidia", "nvidia-build", "nvidia-api"]:
        kwargs.update({
            "api_key": NVIDIA_API_KEY,
        })
    elif provider in ["lm-studio", "lmstudio", "openai-compatible"]:
        kwargs.update({
            "base_url": LM_STUDIO_BASE_URL,
            "api_key": LM_STUDIO_API_KEY if LM_STUDIO_API_KEY else None
        })
    else:
        kwargs["device"] = "cpu"

    return kwargs


def _file_hash(file_path: str) -> str:
    """Compute a fast hash for deduplication (first 8KB + last 8KB + size)."""
    h = hashlib.sha256()
    file_size = os.path.getsize(file_path)
    h.update(str(file_size).encode())

    with open(file_path, "rb") as f:
        h.update(f.read(8192))
        if file_size > 8192:
            f.seek(-8192, 2)
            h.update(f.read(8192))

    return h.hexdigest()[:16]


def _load_single_file(file_path: Path, session_id: str) -> Tuple[str, str, str]:
    """Load a single file and return (file_name, content, file_hash).
    
    Uses robust PDF parsing for PDFs and streaming CSV for large CSVs.
    Falls back to plain text for other formats.
    
    Returns:
        (file_name, content, file_hash) tuple
    """
    file_name = file_path.name
    ext = file_path.suffix.lower()
    fhash = _file_hash(str(file_path))
    content = ""

    try:
        if ext == ".pdf":
            content = load_pdf_robust(str(file_path))

        elif ext in (".csv", ".tsv"):
            # Pass raw CSV text — TextChunker._chunk_csv() expects raw CSV format
            # to detect Q&A structure and do intelligent row-group chunking.
            # Only use streaming for truly massive CSVs (>50MB) that would OOM
            file_size = os.path.getsize(str(file_path))
            if file_size > 50 * 1024 * 1024:
                logger.info(f"Massive CSV ({file_size / 1024 / 1024:.1f}MB), using streaming: {file_name}")
                chunks = []
                for _, text in load_csv_streaming(str(file_path)):
                    chunks.append(text)
                content = "\n\n".join(chunks)
            else:
                # Read raw text — let TextChunker handle CSV parsing natively
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    content = f.read()

        elif ext in (".xlsx", ".xls"):
            try:
                import pandas as pd
                df = pd.read_excel(str(file_path))
                header = " | ".join(str(c) for c in df.columns)
                lines = [f"Columns: {header}"]
                for _, row in df.iterrows():
                    parts = []
                    for col in df.columns:
                        val = row[col]
                        if pd.notna(val):
                            val_str = str(val).strip()
                            if val_str:
                                parts.append(f"{col}: {val_str}")
                    if parts:
                        lines.append(" | ".join(parts))
                content = "\n".join(lines)
            except ImportError:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    content = f.read()

        elif ext in (".json", ".yaml", ".yml", ".md", ".html", ".htm",
                      ".py", ".js", ".java", ".cpp", ".c", ".cs", ".go",
                      ".rs", ".ts", ".jsx", ".tsx", ".txt", ".docx"):
            # For docx, try python-docx first, then unstructured
            if ext == ".docx":
                docx_loaded = False
                
                # Try python-docx first (lightweight, no heavy deps)
                try:
                    import docx
                    doc = docx.Document(str(file_path))
                    
                    parts = []
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(
                                cell.text.strip() for cell in row.cells if cell.text.strip()
                            )
                            if row_text:
                                parts.append(row_text)
                    
                    if parts:
                        content = "\n\n".join(parts)
                        docx_loaded = True
                        logger.info(f"Loaded DOCX with python-docx: {file_name} ({len(parts)} sections)")
                except ImportError:
                    logger.debug(f"python-docx not available, trying unstructured for {file_name}")
                except Exception as e:
                    logger.warning(f"python-docx failed for {file_name}: {e}, trying unstructured")
                
                # Fallback to unstructured if python-docx didn't work
                if not docx_loaded:
                    try:
                        from unstructured.partition.docx import partition_docx
                        elements = partition_docx(filename=str(file_path))
                        content = "\n\n".join(str(el) for el in elements)
                        if content.strip():
                            docx_loaded = True
                    except ImportError:
                        logger.warning(f"Neither python-docx nor unstructured available for DOCX: {file_name}")
                    except Exception as e:
                        logger.warning(f"unstructured also failed for {file_name}: {e}")
                
                if not docx_loaded:
                    logger.error(f"Could not load DOCX: {file_name}. Install python-docx: pip install python-docx")
                    content = ""
            else:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    content = f.read()
        else:
            # Fallback: try reading as text
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()

        # Initialize psycopg2 connection and create document
        doc_id = f"{session_id[:8]}_{fhash}"
        file_size = os.path.getsize(str(file_path))
        
        try:
            conn = psycopg2.connect(**postgres_connect_kwargs(connect_timeout=10))
            with conn.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO documents 
                    (id, name, path, file_type, pages, file_size, hash, loaded_at, status)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (id) DO UPDATE SET
                        name = EXCLUDED.name,
                        path = EXCLUDED.path,
                        file_type = EXCLUDED.file_type,
                        pages = EXCLUDED.pages,
                        file_size = EXCLUDED.file_size,
                        hash = EXCLUDED.hash,
                        loaded_at = EXCLUDED.loaded_at,
                        status = EXCLUDED.status
                """, (
                    doc_id, file_name, str(file_path), ext.lstrip('.'),
                    1, file_size, fhash, datetime.now().isoformat(), "completed" if content.strip() else "failed"
                ))
            conn.commit()
            conn.close()
        except Exception as db_e:
            logger.error(f"Failed to save document {file_name} to database: {db_e}")

        return (file_name, content, fhash)

    except Exception as e:
        logger.error(f"Failed to load {file_name}: {e}")
        # Log failure state to Database if possible
        try:
            conn = psycopg2.connect(**postgres_connect_kwargs(connect_timeout=10))
            with conn.cursor() as cursor:
                 cursor.execute("""
                    INSERT INTO documents 
                    (id, name, path, file_type, pages, file_size, hash, loaded_at, status)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (id) DO UPDATE SET status = EXCLUDED.status
                 """, (
                    f"failed_{int(time.time())}_{file_name}", file_name, str(file_path), file_path.suffix.lstrip('.'),
                    1, 0, "failed_hash", datetime.now().isoformat(), "failed"
                 ))
            conn.commit()
            conn.close()
        except Exception as db_err:
            logger.error(f"Failed to log document error to DB: {db_err}")
            
        raise


def _embed_with_retry(
    embedding_service,
    texts: List[str],
    max_retries: int = MAX_RETRIES,
    base_delay: float = RETRY_BASE_DELAY
) -> list:
    """Embed texts with exponential backoff retry on failure.
    
    Args:
        embedding_service: The embedding service instance
        texts: List of texts to embed
        max_retries: Maximum retry attempts
        base_delay: Base delay in seconds (doubles each retry)
        
    Returns:
        List of embedding vectors
    """
    last_error = None

    for attempt in range(max_retries + 1):
        try:
            embeddings = embedding_service.embed_batch(texts)
            return embeddings
        except Exception as e:
            last_error = e
            if attempt < max_retries:
                delay = base_delay * (2 ** attempt)
                logger.warning(
                    f"Embedding failed (attempt {attempt + 1}/{max_retries + 1}), "
                    f"retrying in {delay:.1f}s: {e}"
                )
                time.sleep(delay)
            else:
                logger.error(f"Embedding failed after {max_retries + 1} attempts: {e}")

    raise last_error


def ingest_documents_session(
    session_id: str,
    documents_dir: Path,
    chunks_dir: Path,
    vector_store_dir: Path,
) -> int:
    """Ingest all documents in a session directory with robust error handling.
    
    Pipeline:
    1. Discover files in documents_dir
    2. Load files in parallel (ThreadPoolExecutor)
    3. Chunk each document using TextChunker (extension-aware routing)
    4. Embed chunks in batches with retry
    5. Store in FAISS vector store
    6. Save metadata
    
    Args:
        session_id: Session identifier
        documents_dir: Path to session's documents directory
        chunks_dir: Path to session's chunks directory 
        vector_store_dir: Path to session's vector store directory
        
    Returns:
        Total number of chunks created
    """
    logger.info(f"[{session_id[:8]}] Starting session ingestion from {documents_dir}")
    start_time = time.time()
    step_times = {}

    # 1. Discover files
    supported_exts = {
        ".pdf", ".csv", ".tsv", ".xlsx", ".xls",
        ".json", ".yaml", ".yml", ".md", ".html", ".htm",
        ".py", ".js", ".java", ".cpp", ".c", ".cs", ".go",
        ".rs", ".ts", ".jsx", ".tsx", ".txt", ".docx",
        ".ipynb",
    }

    files = [
        f for f in Path(documents_dir).iterdir()
        if f.is_file() and f.suffix.lower() in supported_exts
    ]

    if not files:
        logger.warning(f"[{session_id[:8]}] No supported files found in {documents_dir}")
        return 0

    logger.info(f"[{session_id[:8]}] Found {len(files)} files to process")

    # 2. Load files in parallel (I/O-bound — threads are ideal)
    step2_start = time.time()
    loaded_docs = []
    failed_files = []

    if len(files) == 1:
        # Single file — no need for threading overhead
        try:
            result = _load_single_file(files[0], session_id)
            if result[1].strip():  # has content
                loaded_docs.append(result)
            else:
                logger.warning(f"[{session_id[:8]}] File produced no content: {files[0].name}")
                failed_files.append((files[0].name, "No content extracted"))
        except Exception as e:
            failed_files.append((files[0].name, str(e)))
    else:
        # Multiple files — load in parallel
        max_workers = min(MAX_CONCURRENT_FILES, len(files))
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_file = {
                executor.submit(_load_single_file, f, session_id): f for f in files
            }

            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    result = future.result(timeout=120)  # 2 min timeout per file
                    if result[1].strip():
                        loaded_docs.append(result)
                    else:
                        logger.warning(f"[{session_id[:8]}] No content: {file_path.name}")
                        failed_files.append((file_path.name, "No content extracted"))
                except Exception as e:
                    logger.error(f"[{session_id[:8]}] Failed to load {file_path.name}: {e}")
                    failed_files.append((file_path.name, str(e)))

    if failed_files:
        logger.warning(
            f"[{session_id[:8]}] {len(failed_files)} files failed to load: "
            + ", ".join(f[0] for f in failed_files)
        )

    if not loaded_docs:
        raise RuntimeError(
            f"All {len(files)} files failed to load. "
            f"Errors: {'; '.join(f'{n}: {e}' for n, e in failed_files)}"
        )

    step_times['load'] = time.time() - step2_start
    logger.info(f"Step 1: Streamed {len(loaded_docs)} documents in {step_times['load']:.2f}s")

    # 3. Chunk documents (with structure-aware analysis)
    step3_start = time.time()
    chunker = TextChunker(
        chunk_size=600,
        chunk_overlap=90,
        use_db=True,  # Save chunks to PostgreSQL for query-time retrieval
        chunks_dir=str(chunks_dir),
    )

    all_chunks = []
    chunks_metadata = {}

    for file_name, content, fhash in loaded_docs:
        try:
            doc_id = f"{session_id[:8]}_{fhash}"

            chunks = chunker.chunk_text(text=content, doc_id=doc_id, doc_name=file_name)

            # Save chunks to PostgreSQL for query-time retrieval
            if chunks:
                chunker.save_chunks_batch(chunks)

            for chunk in chunks:
                all_chunks.append(chunk)
                chunks_metadata[chunk.id] = {
                    "id": chunk.id,
                    "text": chunk.text,
                    "document_id": doc_id,
                    "document_name": file_name,
                    "chunk_index": chunk.metadata.chunk_index,
                    "token_count": chunk.metadata.token_count,
                    "strategy": chunk.metadata.strategy,
                    "page_number": chunk.metadata.page_number,
                    "source_url": chunk.metadata.source_url,
                }

            logger.info(f"[{session_id[:8]}] Chunked {file_name}: {len(chunks)} chunks")

        except Exception as e:
            logger.error(f"[{session_id[:8]}] Failed to chunk {file_name}: {e}")
            failed_files.append((file_name, f"Chunking error: {e}"))

    if not all_chunks:
        raise RuntimeError(f"No chunks generated from {len(loaded_docs)} loaded documents")

    step_times['chunk'] = time.time() - step3_start
    logger.info(f"Step 2: Chunked {len(all_chunks)} chunks from {len(loaded_docs)} documents in {step_times['chunk']:.2f}s")

    # 4. Save chunks metadata
    import json
    chunks_dir = Path(chunks_dir)
    chunks_dir.mkdir(parents=True, exist_ok=True)
    meta_path = chunks_dir / "chunks_metadata.json"
    prev_path = chunks_dir / "chunks_metadata_prev.json"

    # Archive the previous metadata before overwriting — one-step rollback if needed.
    if meta_path.exists():
        import shutil
        shutil.copy2(str(meta_path), str(prev_path))
        logger.info(f"[{session_id[:8]}] Archived previous chunks_metadata.json → chunks_metadata_prev.json")

    with open(meta_path, "w") as f:
        json.dump(chunks_metadata, f, indent=2)

    # 5. Embed chunks in batches with retry
    step5_start = time.time()
    logger.info(f"STEP 3: Generating Embeddings for {len(all_chunks)} chunks")
    embedding_service = create_embedding_service(
        model_type=EMBEDDING_PROVIDER,
        **_build_embedding_kwargs()
    )

    dimension = embedding_service.model.dimension
    use_qdrant = VECTOR_BACKEND == "qdrant"
    if use_qdrant:
        logger.info(f"[{session_id[:8]}] Using qdrant backend")
        vector_store = QdrantSessionStore(
            session_id=session_id,
            embedding_dimension=dimension,
        )
    else:
        logger.info(f"[{session_id[:8]}] Using FAISS backend")
        vector_store = FAISSVectorStore(
            dimension=dimension,
            index_type="flat",
            store_path=str(vector_store_dir),
        )

    from src.modules.Retriever import BM25Retriever
    bm25_retriever = BM25Retriever(store_path=str(vector_store_dir))

    # Split into batches upfront
    batches = [
        all_chunks[i : i + EMBEDDING_BATCH_SIZE]
        for i in range(0, len(all_chunks), EMBEDDING_BATCH_SIZE)
    ]
    num_batches = len(batches)

    EMBED_WORKERS = 6           # concurrent API calls
    PARALLEL_THRESHOLD = 5      # use concurrent path when more than this many batches

    def _embed_one_batch(batch_chunks):
        """Embed one batch; returns (batch_chunks, results). Runs in a worker thread."""
        texts = [c.text for c in batch_chunks]
        results = _embed_with_retry(embedding_service, texts)
        return batch_chunks, results

    def _write_vectors(batch_chunks, embedding_results):
        """Write one completed embedding batch to the configured vector backend."""
        vectors = [r.embedding for r in embedding_results]
        metadata_list = [
            {
                "session_id":    session_id,
                "document_id":   chunk.metadata.document_id,
                "chunk_id":      chunk.id,
                "text":          chunk.text,
                "document_name": chunk.metadata.document_name,
                "page_number":   chunk.metadata.page_number,
                "source_url":    chunk.metadata.source_url,
            }
            for chunk in batch_chunks
        ]
        vector_store.add_vectors(
            vectors=vectors,
            metadata=metadata_list,
            ids=[c.id for c in batch_chunks],
        )
        
        bm25_retriever.add_texts(
            chunk_ids=[c.id for c in batch_chunks],
            texts=[c.text for c in batch_chunks]
        )

    total_embedded = 0

    if num_batches > PARALLEL_THRESHOLD:
        # ── Concurrent path: multiple API calls in flight simultaneously ──────
        workers = min(EMBED_WORKERS, num_batches)
        logger.info(
            f"[{session_id[:8]}] {num_batches} batches > threshold {PARALLEL_THRESHOLD} "
            f"— embedding concurrently ({workers} workers)"
        )
        with ThreadPoolExecutor(
            max_workers=workers, thread_name_prefix="embedder"
        ) as pool:
            future_to_idx = {
                pool.submit(_embed_one_batch, b): i
                for i, b in enumerate(batches)
            }
            for future in as_completed(future_to_idx):
                idx = future_to_idx[future]
                try:
                    batch_chunks, embedding_results = future.result()
                    _write_vectors(batch_chunks, embedding_results)
                    total_embedded += len(batch_chunks)
                    logger.debug(
                        f"[{session_id[:8]}] Embedded batch {idx + 1}/{num_batches}: "
                        f"{total_embedded}/{len(all_chunks)} chunks"
                    )
                except Exception as e:
                    logger.error(
                        f"[{session_id[:8]}] Embedding batch {idx + 1} failed: {e}\n"
                        f"{traceback.format_exc()}"
                    )
                    failed_files.append(("embedding_batch", str(e)))
    else:
        # ── Sequential path: used for small chunk counts ───────────────────────
        logger.info(
            f"[{session_id[:8]}] {num_batches} batches <= threshold {PARALLEL_THRESHOLD} "
            f"— embedding sequentially"
        )
        for idx, batch in enumerate(batches):
            try:
                batch_chunks, embedding_results = _embed_one_batch(batch)
                _write_vectors(batch_chunks, embedding_results)
                total_embedded += len(batch)
                logger.debug(
                    f"[{session_id[:8]}] Embedded batch {idx + 1}/{num_batches}: "
                    f"{total_embedded}/{len(all_chunks)} chunks"
                )
            except Exception as e:
                logger.error(
                    f"[{session_id[:8]}] Embedding batch {idx + 1} failed: {e}\n"
                    f"{traceback.format_exc()}"
                )
                failed_files.append(("embedding_batch", str(e)))

    step_times['embed'] = time.time() - step5_start
    logger.info(f"Added {total_embedded} vectors to store")
    logger.info(f"Step 3: Generated {total_embedded} embeddings in {step_times['embed']:.2f}s")

    # 6. Save vector store (FAISS requires explicit save; qdrant is persisted on write)
    if hasattr(vector_store, "save"):
        vector_store.save(str(vector_store_dir))
        logger.info(f"Vector store saved with {total_embedded} vectors")
    else:
        logger.info(f"qdrant points upserted: {total_embedded}")

    bm25_retriever.build()
    bm25_retriever.save()
    logger.info("BM25 sparse index built and saved")

    elapsed = time.time() - start_time
    docs_per_sec = len(loaded_docs) / elapsed if elapsed > 0 else 0

    # 7. Background enrichment (optional, non-blocking)
    enable_enrichment = os.environ.get('ENABLE_CHUNK_ENRICHMENT', '').lower() in ('1', 'true', 'yes')
    if enable_enrichment:
        def _run_enrichment():
            try:
                enricher = MetadataEnricher()
                enriched = enricher.enrich_pending_chunks(batch_size=10)
                logger.info(f"[{session_id[:8]}] Background enrichment complete: {enriched} chunks enriched")
            except Exception as e:
                logger.error(f"[{session_id[:8]}] Background enrichment failed: {e}")

        enrichment_thread = threading.Thread(target=_run_enrichment, daemon=True, name="enrichment-worker")
        enrichment_thread.start()
        logger.info(f"[{session_id[:8]}] Background enrichment started for {len(all_chunks)} pending chunks")
    else:
        logger.debug("Chunk enrichment disabled (set ENABLE_CHUNK_ENRICHMENT=1 to enable)")

    # Final summary banner
    logger.info("=" * 80)
    logger.info("PIPELINE BUILD COMPLETE")
    logger.info(
        f"\n"
        f"        Documents: {len(loaded_docs)}\n"
        f"        Chunks: {len(all_chunks)}\n"
        f"        Vectors: {total_embedded}\n"
        f"        Time: {elapsed:.2f}s\n"
        f"        Docs/sec: {docs_per_sec:.2f}"
    )
    logger.info("=" * 80)

    if failed_files:
        logger.warning(
            f"[{session_id[:8]}] Partial failures: "
            + "; ".join(f"{n}: {e}" for n, e in failed_files)
        )

    return total_embedded
